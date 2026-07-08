from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate

from app.database import BASE_DIR


TEMPLATE_PATH = BASE_DIR / "data" / "template_nghidinh30.docx"
CONGVAN_TEMPLATE_PATH = BASE_DIR / "data" / "template_congvan.docx"
OUTPUT_DIR = BASE_DIR / "outputs"


def ensure_default_template() -> Path:
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH

    document = Document()
    set_document_defaults(document)

    # 1. Header table (2 columns)
    header = document.add_table(rows=1, cols=2)
    header.autofit = True
    left = header.cell(0, 0)
    right = header.cell(0, 1)
    
    # Left: Agency Name and Document Code
    left.text = "{{ co_quan }}\n-------------------\nSố: {{ so_ky_hieu }}"
    # Right: National motto and date
    right.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n-----------------------\n{{ dia_danh }}, ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}"

    for cell in (left, right):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    
    # 2. Document type and subject
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_type = title.add_run("{{ loai_van_ban }}\n")
    run_type.bold = True
    run_type.font.size = Pt(14)
    run_subject = title.add_run("{{ trich_yeu }}")
    run_subject.bold = True
    
    document.add_paragraph("")

    # 3. Content body
    document.add_paragraph("{{ noi_dung_ai_viet }}")

    document.add_paragraph("")

    # 4. Footer table
    footer = document.add_table(rows=1, cols=2)
    footer.cell(0, 0).text = "Nơi nhận:\n- Như trên;\n- Lưu: {{ noi_luu }}."
    footer.cell(0, 1).text = "{{ chuc_vu_nguoi_ky }}\n\n\n{{ nguoi_ky }}"
    
    for paragraph in footer.cell(0, 1).paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True

    document.save(TEMPLATE_PATH)
    return TEMPLATE_PATH


def ensure_congvan_template() -> Path:
    CONGVAN_TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    if CONGVAN_TEMPLATE_PATH.exists():
        return CONGVAN_TEMPLATE_PATH

    document = Document()
    set_document_defaults(document)

    header = document.add_table(rows=1, cols=2)
    header.autofit = True
    left = header.cell(0, 0)
    right = header.cell(0, 1)
    left.text = "{{ co_quan }}\n{{ don_vi_soan_thao }}\n-------------------\nSố: {{ so_ky_hieu }}\nV/v {{ trich_yeu }}"
    right.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n-----------------------\n{{ dia_danh }}, ngày {{ ngay }} tháng {{ thang }} năm {{ nam }}"

    for cell in (left, right):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    recipient = document.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recipient.add_run("Kính gửi: {{ don_vi_nhan }}.").bold = True

    document.add_paragraph("{{ doan_mo_dau }}")
    document.add_paragraph("{{ doan_noi_dung_1 }}")
    document.add_paragraph("2. Về việc cử cán bộ tham gia hoạt động: {{ don_vi_soan_thao }} cử {{ so_luong_cb }} đồng chí có tên sau:")

    table = document.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    headers = ["TT", "Họ và tên", "Năm sinh", "Chức vụ", "Ghi chú"]
    for index, title in enumerate(headers):
        run = table.cell(0, index).paragraphs[0].add_run(title)
        run.bold = True
        table.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.cell(1, 0).text = "{%tr for cb in danh_sach_can_bo %}"
    table.cell(2, 0).text = "{{ cb.stt }}"
    table.cell(2, 1).text = "{{ cb.ho_ten }}"
    table.cell(2, 2).text = "{{ cb.nam_sinh }}"
    table.cell(2, 3).text = "{{ cb.chuc_vu }}"
    table.cell(2, 4).text = "{{ cb.ghi_chu }}"
    table.cell(3, 0).text = "{%tr endfor %}"

    document.add_paragraph("{{ doan_ket_thuc }}")

    footer = document.add_table(rows=1, cols=2)
    footer.cell(0, 0).text = "Nơi nhận:\n- Như trên {{ ghi_chu_noi_nhan }};\n- Lưu: {{ noi_luu }}."
    footer.cell(0, 1).text = "{{ chuc_vu_nguoi_ky }}\n\n\n{{ nguoi_ky }}"
    for paragraph in footer.cell(0, 1).paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True

    document.save(CONGVAN_TEMPLATE_PATH)
    return CONGVAN_TEMPLATE_PATH


def set_document_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)


def render_word_document(data: dict) -> Path:
    template_type = data.get("template_type", "default")
    if template_type == "congvan":
        return render_congvan_document(data)

    template_path = ensure_default_template()
    today = datetime.now()
    context = {
        "co_quan": data.get("co_quan", ""),
        "so_ky_hieu": data.get("so_ky_hieu", ""),
        "dia_danh": data.get("dia_danh", ""),
        "ngay": data.get("ngay", f"{today.day:02d}"),
        "thang": data.get("thang", f"{today.month:02d}"),
        "nam": data.get("nam", str(today.year)),
        "loai_van_ban": data.get("loai_van_ban", ""),
        "trich_yeu": data.get("trich_yeu", ""),
        "noi_dung_ai_viet": data.get("noi_dung_ai_viet", ""),
        "noi_luu": data.get("noi_luu", "VT"),
        "chuc_vu_nguoi_ky": data.get("chuc_vu_nguoi_ky", data.get("chuc_vu", "")),
        "nguoi_ky": data.get("nguoi_ky", ""),
    }
    return render_template_to_output(template_path, context, prefix="van-ban")


def render_congvan_document(data: dict) -> Path:
    template_path = ensure_congvan_template()
    today = datetime.now()
    staff = data.get("danh_sach_can_bo", [])

    # Ensure context has all keys used by the Word template.
    context = {
        "co_quan": data.get("co_quan", ""),
        "don_vi_soan_thao": data.get("don_vi_soan_thao", data.get("co_quan", "")),
        "so_ky_hieu": data.get("so_ky_hieu", ""),
        "dia_danh": data.get("dia_danh", ""),
        "ngay": data.get("ngay", f"{today.day:02d}"),
        "thang": data.get("thang", f"{today.month:02d}"),
        "nam": data.get("nam", str(today.year)),
        "trich_yeu": data.get("trich_yeu", ""),
        "don_vi_nhan": data.get("don_vi_nhan", ""),
        "doan_mo_dau": data.get("doan_mo_dau", ""),
        "doan_noi_dung_1": data.get("doan_noi_dung_1", ""),
        "so_luong_cb": str(len(staff)),
        "danh_sach_can_bo": staff,
        "doan_ket_thuc": data.get("doan_ket_thuc", ""),
        "ghi_chu_noi_nhan": data.get("ghi_chu_noi_nhan", ""),
        "noi_luu": data.get("noi_luu", "VT"),
        "chuc_vu_nguoi_ky": data.get("chuc_vu_nguoi_ky", data.get("chuc_vu", "")),
        "nguoi_ky": data.get("nguoi_ky", ""),
        # also used by default template footer/table texts
        "doan_ket_thuc": data.get("doan_ket_thuc", ""),
    }

    try:
        return render_template_to_output(template_path, context, prefix="cong-van")
    except ValueError as exc:
        # Word template might be an older/broken docx (Jinja tag mismatch like missing endfor).
        # Rebuild template and retry once.
        msg = str(exc).lower()
        if "unexpected end of template" in msg or "endfor" in msg:
            template_path = ensure_congvan_template()
            return render_template_to_output(template_path, context, prefix="cong-van")
        raise



def render_template_to_output(template_path: Path, context: dict, prefix: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = DocxTemplate(template_path)
    document.render(context)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    output_path = OUTPUT_DIR / f"{prefix}-{timestamp}-{uuid4().hex[:8]}.docx"
    document.save(output_path)
    return output_path
