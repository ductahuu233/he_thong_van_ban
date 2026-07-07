import pytest
import os
from docxtpl import DocxTemplate
from docx import Document

# Sample data for rendering
sample_data = {
    "title": "Báo cáo cuối cùng",
    "name": "Nguyễn Văn A",
    "items": [
        {"id": 1, "description": "Mục 1", "quantity": 10, "price": 100},
        {"id": 2, "description": "Mục 2", "quantity": 5, "price": 50},
    ]
}

def test_render_docx_template(tmp_path):
    """
    Tests the rendering of a docx template using docxtpl with sample data,
    verifying direct field replacement and table iteration.
    """
    template_path = os.path.join("data", "template_test.docx")
    output_file_name = "rendered_test.docx"
    output_path = tmp_path / output_file_name

    # Dynamically create the test template
    os.makedirs("data", exist_ok=True)
    from docx.shared import Pt
    doc_setup = Document()
    style = doc_setup.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    doc_setup.add_paragraph("{{ title }}")
    doc_setup.add_paragraph("{{ name }}")
    table = doc_setup.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Description", "Quantity", "Price", "Notes"]
    for i, title in enumerate(headers):
        table.cell(0, i).text = title
    table.cell(1, 0).text = "{%tr for item in items %}"
    table.cell(2, 0).text = "{{ item.id }}"
    table.cell(2, 1).text = "{{ item.description }}"
    table.cell(2, 2).text = "{{ item.quantity }}"
    table.cell(2, 3).text = "{{ item.price }}"
    table.cell(2, 4).text = "{{ item.notes or '' }}"
    table.cell(3, 0).text = "{%tr endfor %}"
    doc_setup.save(template_path)

    # Check if the template file exists
    assert os.path.exists(template_path), f"Template file not found at {template_path}"

    # Render the document
    doc = DocxTemplate(template_path)
    doc.render(sample_data)
    doc.save(output_path)

    # Verify the output file exists
    assert os.path.exists(output_path)

    # Open the rendered document and verify its content
    document = Document(output_path)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    document_content = "\n".join(full_text)

    # Verify direct field replacement
    assert sample_data["title"] in document_content
    assert sample_data["name"] in document_content

    # Verify table iteration (check for descriptions of items)
    for item in sample_data["items"]:
        assert item["description"] in document_content
        assert str(item["quantity"]) in document_content
        assert str(item["price"]) in document_content

